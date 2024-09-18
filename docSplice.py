from docx import Document

def split_docx(doc_path, split_point1, split_point2, output1, output2, output3):
    # Load the document
    doc = Document(doc_path)

    # Initialize variables to store the content for the three parts
    part1 = []
    part2 = []
    part3 = []

    # Track which part we're currently filling
    current_part = part1

    # Track whether we've found the split points
    found_first_split = False
    found_second_split = False

    # Loop through the paragraphs in the document
    for para in doc.paragraphs:
        text = para.text

        # Check if the current paragraph contains the first split point
        if split_point1 in text and not found_first_split:
            current_part = part2
            found_first_split = True
        
        # Check if the current paragraph contains the second split point
        elif split_point2 in text and not found_second_split:
            current_part = part3
            found_second_split = True
        
        # Add the paragraph to the appropriate part
        current_part.append(para)

    # Function to write a list of paragraphs to a new document
    def write_part(paragraphs, output_path):
        new_doc = Document()
        for para in paragraphs:
            new_doc.add_paragraph(para.text)
        new_doc.save(output_path)

    # Write the three parts to separate files
    write_part(part1, output1)
    write_part(part2, output2)
    write_part(part3, output3)

# Usage
doc_path = 'Tegus.docx'
split_point1 = 'Managing Partner and Co-Founder of DRIVEN-4'
split_point2 = 'Principal Mechanical Engineer at Zaiput Flow Technologies'
output1 = 'part1.docx'
output2 = 'part2.docx'
output3 = 'part3.docx'

split_docx(doc_path, split_point1, split_point2, output1, output2, output3)